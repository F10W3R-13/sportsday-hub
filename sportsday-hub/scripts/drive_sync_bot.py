# -*- coding: utf-8 -*-
"""
드라이브→로컬 동기화 + 변경 요약 카카오톡 단체방 방송 (작업 스케줄러 매일 22:00).

ZCode 예약 자동화를 대체하는 독립 스크립트 — ZCode 앱 실행 여부와 무관하게
Windows 작업 스케줄러(태스크 GdriveSyncBroadcast)가 직접 실행한다.
18시 봇(KakaoChecklistDigest)과 같은 방식이라 PC가 켜져 있으면 반드시 돈다
(StartWhenAvailable: 놓친 실행은 로그인 즉시 보충, RestartOnFailure: 5분 간격 2회 재시도).

절차: 드라이브 스캔(매니페스트 mtime 비교) → 변경분 다운로드(메모리) → 기계 diff로
방송 메시지 구성 → 카톡 발송 → 성공 시에만 로컬 덮어쓰기 + 매니페스트 갱신 + git 커밋.
발송이 가장 취약한 단계(UI 자동화)라 커밋을 마지막에 둬야 재시도가 멱등하다 —
실패 재시도 시 커밋된 게 없으므로 재탐지·재발송해도 중복이 생기지 않는다.

인증: gdrive MCP가 쓰는 OAuth 리프레시 토큰 재사용(~/.config/mcp-gdrive/).
토큰이 무효화되면 AUTH_FAIL 로그 후 종료(카톡 방송 없음) — ZCode에서 gdrive MCP
도구를 1회 호출해 브라우저 재인증하면 토큰 파일이 갱신돼 봇도 자동 복구된다.

실행: python -X utf8 drive_sync_bot.py
드라이런: SYNC_BOT_DRY=1 — 스캔·diff·메시지 구성까지만(발송·파일 쓰기·커밋 없음).
종료 코드: 0 성공(변동 없음 포함) / 1 실패(로그에 AUTH_FAIL·SEND_FAIL·COMMIT_FAIL).
행사 종료(9/20 18:00) 이후에는 아무 것도 하지 않고 종료.
"""

import csv
import io
import json
import os
import re
import subprocess
import sys
import time
import unicodedata
from collections import OrderedDict
from datetime import datetime, timedelta, timezone
from difflib import SequenceMatcher
from pathlib import Path

import requests

sys.path.insert(0, str(Path(__file__).resolve().parent))
import kakao_group_sender as k  # noqa: E402  (로그·카톡 UI 자동화 재사용)

REPO_ROOT = Path(__file__).resolve().parents[2]
MANIFEST_PATH = REPO_ROOT / ".gdrive-sync-manifest.json"
CREDS_DIR = Path.home() / ".config" / "mcp-gdrive"
KEYS_FILE = CREDS_DIR / "gcp-oauth.keys.json"
TOKEN_FILE = CREDS_DIR / ".gdrive-server-credentials.json"

DRY_RUN = bool(os.environ.get("SYNC_BOT_DRY"))
KST = timezone(timedelta(hours=9))
TOKEN_URL = "https://oauth2.googleapis.com/token"
DRIVE_FILES = "https://www.googleapis.com/drive/v3/files"
LIST_FIELDS = "nextPageToken, files(id, name, mimeType, size, modifiedTime)"
MSG_GAP_SEC = 2.0
WEEKDAYS = "월화수목금토일"

GOOGLE_NATIVE_PREFIX = "application/vnd.google-apps."
EXPORT_MIME = {
    "export-xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "export-docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "export-pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
}
TEXT_EXPORT_MIME = {"spreadsheet": "text/csv", "document": "text/markdown"}
NEW_EXT = {"text-md": ".md", "export-xlsx": ".xlsx", "export-docx": ".docx", "export-pptx": ".pptx"}
# 카톡 1통 팀 표기 — localBase 다음 경로 세그먼트(로컬 폴더명)에서 유추.
TEAM_BY_SEGMENT = [
    ("예산팀", "예산팀"), ("컨텐츠팀", "컨텐츠팀"), ("교환담당팀", "교환담당팀"),
    ("타임라인_인원관리팀", "인원관리팀"), ("25 스포츠데이 참고용 자료", "참고자료"),
]
FOLDER_MAP = {"탐라/인관팀": "타임라인_인원관리팀"}


class AuthFail(Exception):
    pass


class BotError(Exception):
    pass


def nfc(s):
    return unicodedata.normalize("NFC", s)


def parse_mtime(s):
    return datetime.fromisoformat(s.replace("Z", "+00:00"))


# ---------------------------------------------------------------- 인증·드라이브 API

def get_access_token():
    keys = json.loads(KEYS_FILE.read_text(encoding="utf-8"))
    oauth = keys.get("installed") or keys.get("web") or {}
    saved = json.loads(TOKEN_FILE.read_text(encoding="utf-8"))
    if not saved.get("refresh_token") or not oauth.get("client_id"):
        raise AuthFail("리프레시 토큰 또는 OAuth 키 없음")
    res = requests.post(TOKEN_URL, data={
        "grant_type": "refresh_token",
        "refresh_token": saved["refresh_token"],
        "client_id": oauth["client_id"],
        "client_secret": oauth.get("client_secret", ""),
    }, timeout=20)
    if res.status_code != 200:
        raise AuthFail(f"토큰 갱신 실패 http={res.status_code}: {res.text[:150]}")
    return res.json()["access_token"]


def drive_get(token, url, params=None, binary=False):
    res = requests.get(url, headers={"authorization": f"Bearer {token}"},
                       params=params or {}, timeout=60)
    if res.status_code == 401:
        raise AuthFail("401 — 액세스 토큰 거부(리프레시 토큰 무효화 의심)")
    if res.status_code != 200:
        raise BotError(f"드라이브 API 실패 http={res.status_code} {url.split('?')[0]}: {res.text[:150]}")
    return res.content if binary else res.json()


def scan_tree(token, root_id, base, folders):
    """root부터 전체 트리를 재귀 스캔. 신규 하위 폴더는 folders 맵에 등록하고,
    각 파일에는 속한 로컬 폴더 경로(parent_local)를 달아 반환한다."""
    scanned = {}

    def walk(fid, local_dir):
        data = drive_get(token, DRIVE_FILES, params={
            "q": f"'{fid}' in parents and trashed = false",
            "fields": LIST_FIELDS, "pageSize": 200,
        })
        for f in data.get("files", []):
            mime = f.get("mimeType", "")
            if mime == GOOGLE_NATIVE_PREFIX + "folder":
                sub_name = FOLDER_MAP.get(f["name"], f["name"].strip())
                sub_dir = nfc(f"{local_dir}/{sub_name}")
                folders[f["id"]] = sub_dir
                walk(f["id"], sub_dir)
            else:
                scanned[f["id"]] = {"name": f["name"], "mime": mime,
                                    "mtime": f.get("modifiedTime", ""),
                                    "size": int(f.get("size") or 0),
                                    "parent_local": local_dir}
        if data.get("nextPageToken"):  # 200개 초과 폴더는 사실상 없음 — 방어적으로 로그만
            k.log.warning("[sync-bot] 페이지네이션 미처리 폴더 존재(200개 초과)")

    walk(root_id, base)
    return scanned


# ---------------------------------------------------------------- 변경 판정·다운로드

def classify(meta):
    """신규 파일의 kind 결정(기존 파일은 매니페스트 kind 우선). None이면 skipped."""
    if not meta["mime"].startswith(GOOGLE_NATIVE_PREFIX):
        return "as-is"
    short = meta["mime"].removeprefix(GOOGLE_NATIVE_PREFIX)
    if short in TEXT_EXPORT_MIME:
        return "text-md"
    if short == "presentation":
        return "export-pptx"
    return None  # form·drawing 등 — export 불가


def download(token, fid, kind, meta):
    """파일 내용을 메모리로 받는다. text-md는 스냅샷 헤더를 붙여 완성한다."""
    if kind == "as-is":
        return drive_get(token, f"{DRIVE_FILES}/{fid}", params={"alt": "media"}, binary=True)
    if kind == "text-md":
        short = meta["mime"].removeprefix(GOOGLE_NATIVE_PREFIX)
        raw = drive_get(token, f"{DRIVE_FILES}/{fid}/export",
                        params={"mimeType": TEXT_EXPORT_MIME[short]}, binary=True)
        content = (f"```csv\n{raw.decode('utf-8').rstrip()}\n```\n" if short == "spreadsheet"
                   else raw.decode("utf-8"))
        today = datetime.now(KST).strftime("%Y-%m-%d")
        m_kst = parse_mtime(meta["mtime"]).astimezone(KST)
        origin = "구글 시트" if short == "spreadsheet" else "구글 문서"
        head = (f"> 구글 드라이브 동기화 스냅샷 ({today}) — 원본: {origin} 「{meta['name']}」 "
                f"(수정일 {m_kst:%Y-%m-%d %H:%M})\n"
                + ("> 변환: CSV (셀 정렬은 원본 시트 기준)\n" if short == "spreadsheet"
                   else "> 변환: 마크다운 (이미지 생략)\n"))
        return head.encode("utf-8") + b"\n" + content.encode("utf-8")
    return drive_get(token, f"{DRIVE_FILES}/{fid}/export",
                     params={"mimeType": EXPORT_MIME[kind]}, binary=True)


# ---------------------------------------------------------------- diff 엔진
# 결과 규격: OrderedDict {그룹명: [(종류, 항목, 상세)]} — 종류는 변경/신규/제외.

def _opcode_pairs(sm):
    """diff 연산을 (종류, old인덱스, new인덱스) 쌍으로 정규화 — 짝은 변경, 나머지는 제외/신규."""
    for _, i1, i2, j1, j2 in sm.get_opcodes():
        pair = min(i2 - i1, j2 - j1)
        for n in range(pair):
            yield "변경", i1 + n, j1 + n
        for i in range(i1 + pair, i2):
            yield "제외", i, None
        for j in range(j1 + pair, j2):
            yield "신규", None, j


def _fmt_num(v):
    v = str(v).strip()
    m = re.fullmatch(r"[₩원\s]*([\d,\.]+)(원)?", v)
    if m:
        try:
            return f"{float(m.group(1).replace(',', '')):,.0f}"
        except ValueError:
            pass
    return v[:40]


def _looks_like_number(s):
    return bool(re.fullmatch(r"[₩원\s]*[\d,\.]+(원)?", s))


def _clip(s, n=45):
    return re.sub(r"https?://\S+", "(링크)", str(s).strip())[:n]


def _extract_body(text):
    """스냅샷 헤더(>로 시작)와 코드펜스를 걷어 내용만 남긴다."""
    lines = [ln for ln in text.splitlines() if not ln.startswith(">")]
    body = "\n".join(lines).strip()
    m = re.search(r"```(?:csv|markdown)?\s*\n(.*?)```", body, re.S)
    return m.group(1) if m else body


def _csv_rows(text):
    return [r for r in csv.reader(io.StringIO(text)) if any(c.strip() for c in r)]


def _row_key(row):
    cells = [c.strip() for c in row if c.strip()]
    return (" - ".join(cells[:2]) if cells else "(무제)")[:45]


def _track_categories(rows):
    """병합셀 CSV 특성상 카테고리는 해당 구간 첫 행의 A열 — 순회하며 전파한다."""
    cats, cur = [], "기타"
    for r in rows:
        first = next((c.strip() for c in r if c.strip()), "")
        if first and not _looks_like_number(first) and len(first) <= 15:
            cur = first
        cats.append(cur)
    return cats


def diff_csv(old_text, new_text):
    """시트 스냅샷 diff. 행 매칭은 SequenceMatcher, 그룹은 A열(카테고리) 추적."""
    old_rows = [tuple(r) for r in _csv_rows(_extract_body(old_text))]
    new_rows = [tuple(r) for r in _csv_rows(_extract_body(new_text))]
    sm = SequenceMatcher(None, old_rows, new_rows, autojunk=False)
    cats_old, cats_new = _track_categories(old_rows), _track_categories(new_rows)
    groups = OrderedDict()

    def add(cat, entry):
        groups.setdefault(cat, []).append(entry)

    for kind, i, j in _opcode_pairs(sm):
        if kind == "변경":
            detail = " · ".join(
                f"{_fmt_num(a)} → {_fmt_num(b)}"
                for a, b in zip(old_rows[i], new_rows[j])
                if a.strip() != b.strip() and (_looks_like_number(a) or _looks_like_number(b)))[:120]
            add(cats_new[j], (kind, _row_key(new_rows[j]), detail))
        elif kind == "신규":
            add(cats_new[j], (kind, _row_key(new_rows[j]), ""))
        else:
            add(cats_old[i], (kind, _row_key(old_rows[i]), ""))
    return groups


def _norm(s):
    """비교용 정규화 — 마크다운 이스케이프 제거·공백 접기. (표시는 원문 그대로)"""
    return re.sub(r"\s+", " ", re.sub(r"\\([_\+\-*.\[\]()#!])", r"\1", str(s))).strip()


def _disp(s, n=40):
    """표시용 — 이스케이프 제거·마크다운 마커(#, -, >) 제거 후 클립."""
    s = _norm(s).lstrip("#->* ").strip()
    return s[:n]


def _changed_span(old, new):
    """줄 안에서 실제로 바뀐 구간만 '옛값 → 새값'으로 뽑는다(워드 단위)."""
    old_w, new_w = _norm(old).split(), _norm(new).split()
    sm = SequenceMatcher(None, old_w, new_w, autojunk=False)
    olds, news = [], []
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op != "equal":
            olds.extend(old_w[i1:i2])
            news.extend(new_w[j1:j2])
    if not olds and not news:
        return ""
    o = _clip(" ".join(olds), 30) or "(없음)"
    n = _clip(" ".join(news), 30) or "(삭제)"
    return f"{o} → {n}"


def _condense(groups, old_n, new_n):
    """변동이 문서 절반 이상이면 전체 나열 대신 '전면 개정' 요약으로 축약."""
    total = sum(len(v) for v in groups.values())
    if old_n >= 30 and total > old_n * 0.4:
        out = OrderedDict([("전면 개정", [
            ("변경", f"문서 전반 수정 — {old_n}줄 중 {total}줄 변동", "")])])
        adds = [e for v in groups.values() for e in v if e[0] == "신규"][:8]
        if adds:
            out["추가된 내용"] = adds
        return out
    return groups


def diff_lines(old_text, new_text):
    """일반 텍스트/마크다운 라인 diff — 바뀐 구간만 짚고 가장 가까운 제목으로 그룹."""
    old_ls = [ln for ln in _extract_body(old_text).splitlines() if _norm(ln)]
    new_ls = [ln for ln in _extract_body(new_text).splitlines() if _norm(ln)]
    sm = SequenceMatcher(None, [_norm(x) for x in old_ls], [_norm(x) for x in new_ls], autojunk=False)
    groups = OrderedDict()

    def heading_of(lines, idx):
        for ln in reversed(lines[:idx]):
            if re.match(r"^#{1,4}\s|^■\s*", ln.strip()):
                return re.sub(r"^[#■\s]+", "", ln.strip())[:20] or "기타"
        return "기타"

    for kind, i, j in _opcode_pairs(sm):
        target, idx = (new_ls, j) if kind != "제외" else (old_ls, i)
        detail = _changed_span(old_ls[i], new_ls[j]) if kind == "변경" else ""
        groups.setdefault(heading_of(target, idx), []).append((kind, _disp(target[idx]), detail))
    return _condense(groups, len(old_ls), len(new_ls))


def diff_docx(old_bytes, new_bytes):
    import docx
    old_ps = [p.text for p in docx.Document(io.BytesIO(old_bytes)).paragraphs if _norm(p.text)]
    new_ps = [p.text for p in docx.Document(io.BytesIO(new_bytes)).paragraphs if _norm(p.text)]
    sm = SequenceMatcher(None, [_norm(x) for x in old_ps], [_norm(x) for x in new_ps], autojunk=False)
    groups = OrderedDict()

    def heading_of(paras, idx):
        for t in reversed(paras[:idx]):
            t = t.strip()
            if len(t) <= 30 and (re.match(r"^#{1,4}\s|^■\s*", t) or re.fullmatch(r"\[.{1,25}\]", t)):
                return t[:20]
        return "본문"

    for kind, i, j in _opcode_pairs(sm):
        target, idx = (new_ps, j) if kind != "제외" else (old_ps, i)
        detail = _changed_span(old_ps[i], new_ps[j]) if kind == "변경" else ""
        groups.setdefault(heading_of(target, idx), []).append((kind, _disp(target[idx]), detail))
    return _condense(groups, len(old_ps), len(new_ps))


def diff_xlsx(old_bytes, new_bytes):
    import openpyxl
    old_wb = openpyxl.load_workbook(io.BytesIO(old_bytes), data_only=True)
    new_wb = openpyxl.load_workbook(io.BytesIO(new_bytes), data_only=True)
    groups = OrderedDict()
    for title in new_wb.sheetnames:
        ns, os_ = new_wb[title], (old_wb[title] if title in old_wb.sheetnames else None)
        entries = []
        for row in ns.iter_rows():
            for c in row:
                old_v = os_.cell(c.row, c.column).value if os_ else None
                if c.value != old_v and not (c.value in (None, "") and old_v in (None, "")):
                    detail = f"{_fmt_num(old_v)} → {_fmt_num(c.value)}" if os_ else ""
                    entries.append(("신규" if os_ is None else "변경",
                                    f"{c.coordinate} {_fmt_num(c.value)}", detail))
        if entries:
            groups[f"시트 {title}"] = entries[:40]
    return groups


def diff_for(kind, path, old_bytes, new_bytes):
    """kind·확장자에 맞는 diff 엔진 선택. 신규 파일(old 없음)은 내용 요약."""
    is_new = old_bytes is None
    suffix = Path(path).suffix.lower()
    try:
        if is_new:
            if suffix == ".md" or kind == "text-md":
                return _summarize_new_md(new_bytes)
            if suffix == ".docx":
                return _summarize_new_docx(new_bytes)
            if suffix == ".xlsx":
                return _summarize_new_xlsx(new_bytes)
            return OrderedDict([("신규", [("신규", f"{len(new_bytes):,}B", "")])])
        if suffix == ".md" or kind == "text-md":
            old_text = old_bytes.decode("utf-8", errors="replace")
            new_text = new_bytes.decode("utf-8", errors="replace")
            if "변환: CSV" in new_text[:300] or "변환: CSV" in old_text[:300]:
                return diff_csv(old_text, new_text)
            return diff_lines(old_text, new_text)
        if suffix == ".docx":
            return diff_docx(old_bytes, new_bytes)
        if suffix == ".xlsx":
            return diff_xlsx(old_bytes, new_bytes)
    except Exception as exc:  # diff 실패는 방송을 막지 않는다 — 크기 비교로 대체
        k.log.warning("[sync-bot] diff 엔진 실패(%s): %s", path, exc)
    return OrderedDict([("변동", [("변경", f"{len(old_bytes):,}B → {len(new_bytes):,}B", "")])])


def _summarize_new_md(data):
    """신규 md/시트 — 전체 나열 대신 규모+제목 한 줄."""
    text = data.decode("utf-8", errors="replace")
    if "```csv" in text:
        rows = _csv_rows(_extract_body(text))
        title = _disp(_row_key(rows[0])) if rows else ""
        desc = f"{len(rows)}행 구성 (헤더 포함)" + (f" — {title}" if title else "")
    else:
        lines = [ln for ln in _extract_body(text).splitlines() if _norm(ln)]
        title = next((_disp(ln, 30) for ln in lines
                      if re.match(r"^#{1,3}\s", ln.strip())), _disp(lines[0], 30) if lines else "")
        desc = f"{len(lines)}줄 구성" + (f" — {title}" if title else "")
    return OrderedDict([("신규", [("신규", desc, "")])])


def _summarize_new_docx(data):
    import docx
    paras = [p.text for p in docx.Document(io.BytesIO(data)).paragraphs if _norm(p.text)]
    return OrderedDict([("신규", [("신규", f"{len(paras)} 문단 구성 — {_disp(paras[0], 30) if paras else ''}", "")])])


def _summarize_new_xlsx(data):
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    desc = " · ".join(f"{t} {wb[t].max_row}행" for t in wb.sheetnames)
    return OrderedDict([("신규", [("신규", f"시트 {len(wb.sheetnames)}개 — {desc}", "")])])


# ---------------------------------------------------------------- 메시지 구성

def team_of(path):
    for seg in Path(path).parts[1:]:  # 첫 세그먼트는 localBase
        for folder, team in TEAM_BY_SEGMENT:
            if seg == folder:
                return team
    return "기획팀"


def stat_line(groups, is_new):
    first = next(iter(groups.values()), None)
    if "전면 개정" in groups:
        return groups["전면 개정"][0][1]
    if is_new:
        return f"신규 문서 — {first[0][1]}" if first else "신규 문서"
    counts = {"변경": 0, "신규": 0, "제외": 0}
    for entries in groups.values():
        for kind, _, _ in entries:
            counts[kind] += 1
    return " · ".join(f"{v}건 {kk}" for kk, v in counts.items() if v) or "미세 변경"


def _settle_group(groups):
    """'총 지출/합계/잔여/지원/수입' 항목은 마지막 ■ 정산 그룹으로 옮긴다(예산 서식)."""
    settle = []
    for cat in list(groups):
        keep = [e for e in groups[cat] if not re.search(r"총\s*지출|합계|잔여|지원|수입|정산", e[1])]
        settle.extend(e for e in groups[cat] if e not in keep)
        if keep:
            groups[cat] = keep
        else:
            del groups[cat]
    if settle:
        groups["정산"] = settle
    return groups


def render_groups(groups, budget):
    """■ 그룹 → • 불릿 렌더. budget자 내에서 불릿 단위로 채우고 남은 건 '외 N건'으로."""
    out, used = [], 0

    def put(line):
        nonlocal used
        out[-1].append(line)
        used += len(line) + 1

    for cat, entries in _settle_group(groups).items():
        header = f"■ {cat}"
        if used + len(header) > budget:
            break
        out.append([header])
        used += len(header) + 1
        shown = 0
        for kind, item, detail in entries:
            tag = {"신규": " (+)", "제외": " (−)"}.get(kind, "")
            bullet = f"• {item}{tag}" + (f": {detail[:70]}" if detail else "")
            if used + len(bullet) > budget - 15:  # '외 N건' 여유 확보
                break
            put(bullet)
            shown += 1
        rest = len(entries) - shown
        if rest > 0:
            put(f"… 외 {rest}건")
            break  # 예산 소진
    return "\n".join("\n".join(block) for block in out)


def compose_messages(changes):
    """확정 서식(2026-08-27) 준수. changes: [{path,display,team,mtime_kst,is_new,groups,stat}]"""
    today = datetime.now(KST)
    if not changes:
        return [f"📂 [드라이브 문서 업데이트] {today.month}/{today.day}({WEEKDAYS[today.weekday()]})\n\n"
                "오늘 드라이브 변경 없습니다.\n문서가 올라오면 이 방으로 안내드릴게요."]

    earliest = min(c["mtime_kst"].date() for c in changes)
    when = "오늘" if earliest == today.date() else f"{earliest.month}/{earliest.day} 이후"
    lines = [f"📂 [드라이브 문서 업데이트] {today.month}/{today.day}({WEEKDAYS[today.weekday()]})",
             "", f"{when} 드라이브에서 갱신된 문서 {len(changes)}건"]
    for c in changes:
        new_tag = " · 신규" if c["is_new"] else ""
        lines.append(f"\n• {c['display']} ({c['team']} · {c['mtime_kst']:%H:%M} 수정{new_tag})")
        lines.append(f"  └ {c['stat']}")
    lines.append("\n핵심 변화는 다음 메시지에 →")
    msg1 = "\n".join(lines)
    if len(msg1) > 300:
        msg1 = msg1[:297].rstrip() + "…"

    per_budget = max(250, 1500 // len(changes) - 40)
    parts = []
    for c in changes:
        body = render_groups(c["groups"], per_budget)
        header = f"[핵심 변화] {c['display']} {c['mtime_kst'].month}/{c['mtime_kst'].day} 개정분"
        parts.append(f"{header}\n\n{body}" if body else f"{header}\n\n(미세 변경 — 원문 참고)")
    msg2 = "\n\n".join(parts)
    if len(msg2) > 1500:
        msg2 = msg2[:1497].rstrip() + "…"
    return [msg1, msg2]


# ---------------------------------------------------------------- 커밋

def git(*args):
    r = subprocess.run(["git", "-C", str(REPO_ROOT), *args], capture_output=True,
                       text=True, encoding="utf-8", errors="replace")
    if r.returncode != 0:
        raise BotError(f"git {args[0]} 실패: {(r.stderr or r.stdout).strip()[:300]}")
    return r.stdout.strip()


# ---------------------------------------------------------------- 메인

def main():
    if datetime.now() >= k.BOT_END:
        k.log.info("[sync-bot] 행사 종료(9/20 18:00) 이후 — 실행 안 함")
        return
    k.log.info("[sync-bot] 실행 시작: dry=%s", DRY_RUN)
    token = get_access_token()
    k.log.info("[sync-bot] 인증 완료")

    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    folders = manifest.setdefault("folders", {})
    known_folders = set(folders)
    base = manifest.get("localBase", "26-2 Sports Day")
    scanned = scan_tree(token, manifest["rootFolderId"], base, folders)
    new_folders = {fid: p for fid, p in folders.items() if fid not in known_folders}
    k.log.info("[sync-bot] 스캔 완료: 파일 %d개 · 폴더 %d개(신규 %d)",
               len(scanned), len(folders), len(new_folders))

    files_map = manifest["files"]
    skipped = manifest.setdefault("skipped", {})
    changed, deleted = [], []

    for fid, meta in scanned.items():
        entry = files_map.get(fid)
        if entry is None:
            kind = classify(meta)
            if kind is None:
                skipped[fid] = {"name": meta["name"], "reason": "구글 폼 등 — export 불가"}
                continue
            suffix = NEW_EXT.get(kind, Path(meta["name"]).suffix)
            stem = meta["name"] if kind == "text-md" else Path(meta["name"]).stem
            rel = nfc(f"{meta['parent_local']}/{stem}{suffix}")
            if any(e.get("path") == rel for e in files_map.values()):
                rel = nfc(f"{meta['parent_local']}/{stem} (신규){suffix}")
            entry = files_map[fid] = {"path": rel, "name": meta["name"], "kind": kind}
            entry.update(mtime=meta["mtime"], size=meta["size"])
            changed.append(_prep_change(token, fid, entry, meta, is_new=True))
        elif parse_mtime(meta["mtime"]) > parse_mtime(entry["mtime"]):
            entry.update(mtime=meta["mtime"], size=meta["size"])
            changed.append(_prep_change(token, fid, entry, meta, is_new=False))
    for fid, entry in files_map.items():
        if fid not in scanned and "삭제됨" not in entry.get("note", ""):
            note = f"드라이브에서 삭제됨 ({datetime.now(KST):%m/%d} 확인)"
            entry["note"] = note
            deleted.append(f"{entry['path']} — {note}")

    if deleted:
        k.log.info("[sync-bot] 드라이브 삭제 확인: %s", "; ".join(deleted))
    messages = compose_messages([
        {"path": c["path"], "display": Path(c["path"]).stem, "team": team_of(c["path"]),
         "mtime_kst": c["mtime_kst"], "is_new": c["is_new"], "groups": c["groups"],
         "stat": c["stat"]}
        for c in changed])
    k.log.info("[sync-bot] 변동 %d건 — 메시지 %d통", len(changed), len(messages))
    if DRY_RUN:
        for i, m in enumerate(messages, 1):
            k.log.info("[sync-bot][dry] %d/%d통: %d자 | %s", i, len(messages), len(m), m.splitlines()[0])
            k.log.info("[sync-bot][dry] 본문:\n%s", m)
        return

    # 발송(취약 단계) → 성공 시에만 파일·매니페스트·커밋
    try:
        k.ensure_kakao_running()
        room = k.open_room_with_retry(k.ROOM_NAME)
        for i, text in enumerate(messages, 1):
            k.send_message(room, text)
            k.log.info("[sync-bot] 발송 %d/%d통: %d자 | %s", i, len(messages), len(text), text.splitlines()[0])
            if i < len(messages):
                time.sleep(MSG_GAP_SEC)
    except Exception as exc:
        k.log.exception("[sync-bot] 카톡 발송 실패 — 커밋하지 않음(재시도 시 재발송됨)")
        print(f"SEND_FAIL: {type(exc).__name__}: {exc}", file=sys.stderr)
        sys.exit(1)

    try:
        paths = []
        for c in changed:
            local = REPO_ROOT / c["path"]
            local.parent.mkdir(parents=True, exist_ok=True)
            local.write_bytes(c["data"])
            paths.append(c["path"])
        manifest["lastSyncedAt"] = datetime.now(KST).isoformat(timespec="seconds")
        MANIFEST_PATH.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
                                 encoding="utf-8")
        subject = f"sync: 드라이브 → 로컬 스냅샷 ({len(changed)}건 갱신)"
        body_lines = [f"- {c['path']}: {c['stat']}" for c in changed]
        body_lines += [f"- 삭제 확인: {d}" for d in deleted]
        body_lines += [f"- 신규 폴더 등록: {p}" for p in new_folders.values()]
        git("add", "--", *paths, ".gdrive-sync-manifest.json")
        git("commit", "-m", subject, "-m", "\n".join(body_lines))
        k.log.info("[sync-bot] 커밋 완료: %s", subject)
    except Exception as exc:
        k.log.exception("[sync-bot] 커밋 실패 — 발송은 완료됨. 매니페스트가 갱신 안 돼 "
                        "다음 실행 때 같은 변경이 재방송될 수 있음(수동 확인 필요)")
        print(f"COMMIT_FAIL: {type(exc).__name__}: {exc}", file=sys.stderr)
        sys.exit(1)
    k.log.info("[sync-bot] 실행 완료")


def _prep_change(token, fid, entry, meta, is_new):
    local = REPO_ROOT / entry["path"]
    old_bytes = local.read_bytes() if (not is_new and local.exists()) else None
    data = download(token, fid, entry["kind"], meta)
    groups = diff_for(entry["kind"], entry["path"], old_bytes, data)
    return {"path": entry["path"], "data": data, "groups": groups,
            "mtime_kst": parse_mtime(meta["mtime"]).astimezone(KST),
            "is_new": is_new, "stat": stat_line(groups, is_new)}


if __name__ == "__main__":
    try:
        main()
    except AuthFail as exc:
        k.log.error("[sync-bot] AUTH_FAIL — gdrive MCP 재인증 필요(ZCode에서 MCP 도구 1회 호출): %s", exc)
        print(f"AUTH_FAIL: {exc}", file=sys.stderr)
        sys.exit(1)
    except BotError as exc:
        k.log.error("[sync-bot] 실패: %s", exc)
        print(f"BOT_FAIL: {exc}", file=sys.stderr)
        sys.exit(1)
